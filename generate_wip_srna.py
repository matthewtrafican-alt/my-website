"""
Generates WIP-SRNA-001 v3.0
Workforce Intelligence Platform — System Requirements & Notional Architecture
Supersedes v2.0. Adds enterprise scale architecture (50K FTEs, 12 functions,
9 programmes, 3 portfolios), Functional / Program dual-view, four temporal
dimensions (time in role / level / programme / acquisition lifecycle), resource
allocation matrix (Function x Programme), phase experience gap analysis, and
AI-integrated prompt library (280 diagnostic prompts).
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading(text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold=False, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p

def action_box(ref, text):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'FFF3CD')
    p    = cell.paragraphs[0]
    p.clear()
    run  = p.add_run(f'ACTION REQUIRED — {ref}')
    run.bold = True
    run.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
    run.font.size = Pt(9)
    cell.add_paragraph(text).paragraph_format.space_after = Pt(3)
    op = cell.add_paragraph()
    op.add_run('Owner: ________________________    Date: ________________________    Status: open / in progress / resolved').font.size = Pt(8)
    doc.add_paragraph()

def req_table(rows_data):
    headers = ['ID', 'Requirement', 'Priority', 'Rationale']
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1F3864')
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
    moscow_colors = {'Must': 'C6EFCE', 'Should': 'FFEB9C', 'Could': 'DDEBF7', "Won't": 'FCE4D6'}
    for row_data in rows_data:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        moscow = row_data[2]
        set_cell_bg(row[2], moscow_colors.get(moscow, 'FFFFFF'))
    doc.add_paragraph()

def simple_table(headers, rows, header_color='1F3864'):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_color)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
    for row_data in rows:
        row = tbl.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

def road_box(pillar, color_hex, definition, outcomes, indicators):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(pillar)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    cell.add_paragraph(definition)
    cell.add_paragraph(f'Outcomes: {outcomes}')
    cell.add_paragraph(f'Leading indicators: {indicators}')
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('S Y S T E M   R E Q U I R E M E N T S   &   N O T I O N A L\nA R C H I T E C T U R E')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
run.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Workforce Intelligence Platform')
run2.font.size = Pt(26)
run2.bold = True
run2.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('AI-Powered Program Delivery Risk Management\n& Strategic Workforce Intelligence')
run3.font.size = Pt(13)
run3.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('ROAD Framework · VUCA Prime · Theory of Constraints · Donella Meadows')
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
run4.italic = True

doc.add_paragraph()

meta = doc.add_table(rows=5, cols=4)
meta.style = 'Table Grid'
meta_data = [
    ['Organisation',  'AeroDefend Group (Illustrative)', 'Version', '3.0 — DRAFT'],
    ['Classification','Internal — Confidential',          'Status',  'Pending IT & CPO Review'],
    ['Prepared by',   'HR Transformation Initiative',     'Prepared for', 'IT VP, HR VP, CPO & Programme Directors'],
    ['Document ref',  'WIP-SRNA-001',                     'Supersedes', 'WIP-SRNA-001 v2.0'],
    ['Date',          '30 May 2026',                      'Base document', 'H2R-SRNA-001 v1.0'],
]
for r, row_data in enumerate(meta_data):
    cells = meta.rows[r].cells
    for c, val in enumerate(row_data):
        cells[c].text = val
        if c % 2 == 0:
            set_cell_bg(cells[c], 'D6E4F0')
            for para in cells[c].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        else:
            for para in cells[c].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

doc.add_paragraph()

info_tbl = doc.add_table(rows=1, cols=1)
info_tbl.style = 'Table Grid'
info_cell = info_tbl.rows[0].cells[0]
set_cell_bg(info_cell, 'D6E4F0')
info_p = info_cell.paragraphs[0]
info_p.clear()
run_info = info_p.add_run('Relationship to H2R-SRNA-001 and WIP-SRNA-001 v1.0')
run_info.bold = True
run_info.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
info_cell.add_paragraph(
    'This document (v2.0) supersedes WIP-SRNA-001 v1.0 and directly extends H2R-SRNA-001 v1.0. '
    'All requirements from both predecessor documents are carried forward and supplemented with: '
    'the ROAD strategic framework (Retain / Optimize / Acquire / Develop); VUCA Prime design '
    'philosophy; Theory of Constraints applied to talent as a product; program execution risk '
    'management as the primary business driver; attrition classification by program impact '
    '(replacing any 9-box based classification); turnover contagion modelling at team level; '
    'proactive action plans with return on impact; three-tier ownership model; and a phased '
    'data architecture from flat files through to live IT integration.'
)
doc.add_paragraph()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. DOCUMENT CONTROL
# ══════════════════════════════════════════════════════════════════════════════
heading('1. Document Control & Revision History', 1)
simple_table(
    ['Version', 'Date', 'Author', 'Changes'],
    [
        ['1.0', '29/05/2026', 'HR Transformation Team', 'Initial draft — multi-horizon workforce planning and competency intelligence'],
        ['2.0', '30/05/2026', 'HR Transformation Team', 'Major revision — ROAD framework, VUCA Prime, Theory of Constraints, program execution risk as primary driver, attrition by program impact, turnover contagion, return on impact action plans, three-tier ownership model, flat-file phasing architecture'],
        ['3.0', '30/05/2026', 'HR Transformation Team', 'Enterprise scale — 50K FTEs, 12 functional organisations, 9 programmes, 3 portfolios. Functional/Program dual-view with scope selector. Four temporal dimensions (time in role, time in level, time on programme, acquisition lifecycle experience). Resource allocation matrix (Function x Programme with headcount, FTE%, flight risk). Phase experience gap analysis. AI-integrated prompt library (280 diagnostic prompts). Updated synthetic data schema v2.'],
    ]
)
heading('1.1 Document Approvals Required', 2)
simple_table(
    ['Role', 'Name', 'Date', 'Signature / Approval'],
    [
        ['IT Vice President', '', '', ''],
        ['HR Vice President', '', '', ''],
        ['Chief People Officer', '', '', ''],
        ['Programme Directors (for Programme Delivery Risk framing)', '', '', ''],
        ['IT Architecture Lead', '', '', ''],
        ['Data Engineering Lead', '', '', ''],
        ['CISO / Information Security', '', '', ''],
        ['Legal / Compliance', '', '', ''],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading('2. Executive Summary', 1)
body(
    'The AeroDefend Group Workforce Intelligence Platform is an analytics and intelligence overlay '
    'that sits above existing HR systems (primarily SAP SuccessFactors and SAP HCM) without '
    'replacing them. It addresses a specific business problem: on-time programme delivery is at '
    'risk when the right talent is not in the right place, at the right time, with the right '
    'skills, at the right cost. Talent is treated as a product in a supply chain. When that supply '
    'chain has a bottleneck, programmes slip, customers are disappointed, and commercial '
    'commitments are at risk. The platform makes workforce risk visible before it becomes a '
    'delivery crisis.'
)
body(
    'This platform is not a performance management tool, not a 9-box succession grid, and not a '
    'case management system. Case management stays in SAP. The platform is the intelligence layer '
    'that generates insight from SAP data and turns it into proactive, prioritised action plans '
    'for the people who can act on them — Line Managers, HR Business Partners, and Leaders — '
    'each seeing the information relevant to their role.'
)
body(
    'Version 2.0 introduces the ROAD strategic framework (Retain, Optimize, Acquire, Develop) as '
    'the organisational architecture for all workforce actions. It introduces VUCA Prime '
    '(Vision, Understanding, Clarity, Agility) as the platform design philosophy — the platform '
    'is the antidote to the volatility, uncertainty, complexity, and ambiguity of the defence '
    'and aerospace operating environment. And it introduces the Theory of Constraints applied to '
    'human capital: the platform identifies the binding constraint in the talent supply chain '
    'and directs attention to the one thing that, if resolved, would have the greatest impact '
    'on programme delivery.'
)

heading('2.1 Primary Business Driver — On-Time Programme Delivery', 2)
body(
    'The primary output metric of this platform is delivery confidence. Every workforce action '
    'recommended by the platform is anchored to a specific programme, a specific milestone, '
    'and a specific impact if no action is taken. This framing is deliberate: HR interventions '
    'are most effective when they are connected to the business outcome they are protecting.'
)
body(
    'The "five rights of talent" framework defines the standard: Right talent — Right place — '
    'Right time — Right skills — Right cost. When any of the five rights is breached, the '
    'platform identifies it, explains why it matters in plain English, and recommends a specific '
    'action with a calculated return on impact.'
)
simple_table(
    ['If this goes wrong...', 'The business impact is...', 'ROAD response is...'],
    [
        ['Right talent not in place (key person leaves)', 'Programme critical path exposed; milestone at risk', 'Retain — flight risk intervention before departure'],
        ['Wrong place (talent allocated suboptimally)', 'Programme bottleneck; underutilisation elsewhere', 'Optimize — redeployment and org design action'],
        ['Wrong time (hired too late — behind lead time)', 'Programme milestone missed; contract penalty risk', 'Acquire — "Ahead of Ready" pipeline acceleration'],
        ['Wrong skills (competency gap on critical role)', 'Delivery quality risk; rework and delay', 'Develop — targeted upskilling with milestone linkage'],
        ['Wrong cost (talent priced out of market)', 'Retention failure; external hire premium cost', 'Retain + Acquire — compensation review and pipeline'],
    ]
)

heading('2.2 Platform Capabilities at a Glance', 2)
simple_table(
    ['Capability', 'Description', 'ROAD Pillar', 'Planning Horizon'],
    [
        ['Program Delivery Risk View', 'Talent supply vs demand by programme, "Ahead of Ready" status, milestone risk', 'All ROAD', 'Tactical (0-12mo)'],
        ['Retain — Flight Risk & Attrition', 'Flight risk scoring, regrettable attrition classification by programme impact, turnover contagion', 'Retain', 'Tactical'],
        ['Optimize — Org Performance', 'STAR Model (Galbraith), span of control, org layer design, productivity indicators', 'Optimize', 'Operational'],
        ['Acquire — Talent Pipeline', '"Ahead of Ready" lead time management, TOC bottleneck identification, H2R case risk scoring', 'Acquire', 'Tactical + Execution'],
        ['Develop — Capability Building', 'Competency gap analysis, succession depth, L&D investment vs gap closure, career pathways', 'Develop', 'Operational'],
        ['Proactive Action Plans', 'Driving indicator-based action recommendations with return on impact (cost of action vs inaction)', 'All ROAD', 'Execution'],
        ['H2R Process Intelligence', 'Process mining, SLA management, compliance monitoring (inherited from H2R-SRNA-001)', 'Acquire + Retain', 'Execution'],
        ['Strategic Workforce Planning', '3-year capability gap analysis, scenario modelling, build/buy/borrow/bot decisions', 'All ROAD', 'Strategic (3-5yr)'],
        ['Competency Intelligence', 'AI-generated competency models grounded in O*NET, NICE, INCOSE SECF, R&M BoK', 'Develop', 'Operational'],
        ['Cascade Engine', 'Bidirectional flow: strategy to execution and data to forecast', 'All ROAD', 'All horizons'],
        ['AI Agent', 'Claude API-powered workforce intelligence agent grounded in real-time data via ChromaDB RAG', 'All ROAD', 'All horizons'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. SCOPE & BUSINESS CONTEXT
# ══════════════════════════════════════════════════════════════════════════════
heading('3. Scope & Business Context', 1)
heading('3.1 What This Platform Is — And Is Not', 2)
body(
    'This platform is an analytics and intelligence overlay. It reads data from SAP and other '
    'sources, processes it, and surfaces insights and action recommendations. It does NOT replace '
    'SAP case management. It does NOT write back to SAP. HR Business Partners use SAP for '
    'execution; they use this platform to decide what to execute.'
)
simple_table(
    ['What the Platform Does', 'What the Platform Does NOT Do'],
    [
        ['Reads H2R event logs from SAP SuccessFactors', 'Create or update SAP cases (case management stays in SAP)'],
        ['Scores flight risk based on engagement, pay, tenure signals', 'Replace performance management — performance stays in SAP'],
        ['Classifies attrition as regrettable or non-regrettable by programme impact', 'Operate a 9-box succession grid (development is separate from performance)'],
        ['Generates action plans with return on impact calculations', 'Automate HR decisions — all decisions are made by humans'],
        ['Shows programme delivery risk from a talent perspective', 'Replace project management systems — EVM stays in programme tools'],
        ['Surfaces the binding constraint in the talent supply chain', 'Remove the constraint autonomously — that requires human intervention'],
    ]
)

heading('3.2 The Four-Horizon Planning Architecture', 2)
simple_table(
    ['Layer', 'Horizon', 'Primary Question', 'Planning Rhythm', 'Primary Output'],
    [
        ['Strategic Workforce Planning', '3-5 years', 'What workforce shape and capability do we need to execute business strategy?', 'Annual + trigger events', 'Strategic capability plan, build/buy/borrow decisions'],
        ['Operational Workforce Planning', '12-36 months', 'How many people, in what roles, with what competencies, and when?', 'Annual budget cycle + quarterly review', 'Headcount plan, L&D investment plan, succession slate'],
        ['Tactical Workforce Planning', '0-12 months', 'Who do we hire, redeploy, develop, or retain — and in what sequence?', 'Monthly + sprint cadence', 'ROAD action plans, hiring plan, retention actions'],
        ['H2R Execution', 'Daily/weekly', 'Are our HR processes running on time and in compliance?', 'Daily pipeline, real-time alerts', 'Process risk scores, SLA breach reports, compliance alerts'],
    ]
)

heading('3.3 In Scope — Phase 1 (Months 1–6)', 2)
bullet('All capabilities from H2R-SRNA-001 Phase 1 — carried forward in full')
bullet('ROAD Framework implementation: Retain, Optimize, Acquire, Develop as the navigation and action backbone')
bullet('Program Delivery Risk view: talent supply vs demand by programme, "Ahead of Ready" status, milestone risk')
bullet('Attrition Intelligence: classification by programme execution impact (regrettable/non-regrettable/neutral), NOT 9-box')
bullet('Turnover Contagion: team-level risk modelling when regrettable departures occur')
bullet('Proactive Action Plans: driving indicator-based recommendations with return on impact')
bullet('Three-tier ownership model: Action Owner (Line Manager) / Business Partner (HR BP) / Leader (HR VP/CPO) views')
bullet('STAR Model (Galbraith): Strategy, Structure, Processes, Rewards, People scored at individual/team/org level')
bullet('Theory of Constraints: binding constraint identification in talent supply chain with lead time analysis')
bullet('Flat files as Phase 1 data architecture — CSV/JSON from SAP export, pre-processed into vector stores')
bullet('Framework Backbone ingestion: O*NET, NICE, INCOSE SECF, R&M BoK, Bloom\'s Taxonomy')
bullet('Strategic Workforce Planning module: 3-year capability gap analysis, scenario modelling')
bullet('Competency Intelligence: AI-generated competency models grounded in framework backbone')
bullet('Cascade Engine: bidirectional planning layer connector')
bullet('AI Agent: Claude API with RAG over ChromaDB vector store')
bullet('Static demo deployment: GitHub Pages (no backend required for demonstration)')

heading('3.4 Explicitly Out of Scope — Phase 1', 2)
bullet('Direct write-back to SAP (read-only integration in Phase 1)')
bullet('Real-time SAP event streaming (batch/nightly ingestion in Phase 1; event-driven in Phase 3)')
bullet('Mobile application')
bullet('AFECD military classification integration (Phase 2)')
bullet('External talent market data (Lightcast, LinkedIn — Phase 2)')
bullet('Automated workflow execution in SAP (Phase 2)')
bullet('Full SIOP-compliant structured job analysis data collection (Phase 2 — framework seeded in Phase 1)')
bullet('9-box grid or any tool blending development with talent performance management')

# ══════════════════════════════════════════════════════════════════════════════
# 4. DESIGN PHILOSOPHY & STRATEGIC FRAMEWORK
# ══════════════════════════════════════════════════════════════════════════════
heading('4. Design Philosophy & Strategic Framework', 1)
body(
    'The Workforce Intelligence Platform is built on four interlocking design principles. '
    'These are not branding choices — they directly shape what the platform measures, '
    'what it surfaces, and how it recommends actions. Understanding these principles is '
    'essential for anyone designing, building, or using the platform.'
)

heading('4.1 VUCA Prime — The Platform as the Antidote', 2)
body(
    'Defence and aerospace organisations operate in a VUCA environment: Volatility '
    '(programme scope changes without warning), Uncertainty (technology risk, competitor moves), '
    'Complexity (multi-programme talent allocation, clearance dependencies), Ambiguity '
    '(unclear future capability requirements). The response to VUCA is not more dashboards — '
    'it is VUCA Prime: Vision, Understanding, Clarity, Agility.'
)
simple_table(
    ['VUCA Challenge', 'VUCA Prime Response', 'Platform Mechanism'],
    [
        ['Volatility — environment changes faster than plans', 'Vision — know where you are going and why', '3-year workforce plan, programme milestone linkage, strategic ROAD targets'],
        ['Uncertainty — cannot predict what will happen', 'Understanding — know why things are happening now', 'Driving indicator analysis, causal narratives in plain English, root cause identification'],
        ['Complexity — too many interacting variables', 'Clarity — know what to do about it', 'Prioritised action plans, single recommended action per risk, owner assigned per action'],
        ['Ambiguity — information is incomplete or unclear', 'Agility — move faster than the problem', 'Real-time data pipeline, ROAD action routing, proactive "Ahead of Ready" lead time management'],
    ]
)

heading('4.2 ROAD Framework — The Workforce Strategy Architecture', 2)
body(
    'ROAD is the strategic framework that organises all workforce actions into four '
    'mutually reinforcing pillars. Every action plan generated by the platform is tagged '
    'to a ROAD pillar. Every metric is assigned to a ROAD category. The framework ensures '
    'workforce decisions are balanced — preventing organisations from over-indexing on '
    'any single dimension (e.g. only hiring when they should also be retaining or developing).'
)
simple_table(
    ['Pillar', 'Full Name', 'Core Outcomes', 'Stabilises Against', 'Primary Indicators'],
    [
        ['R', 'Retain', 'Stabilise critical skills and foster a positive work environment', 'Regrettable attrition, contagion spread, clearance loss, knowledge walk-out', 'Flight risk score, compa-ratio, engagement score, attrition rate (regrettable), contagion index'],
        ['O', 'Optimize', 'Accelerate organisational performance and maximise productivity', 'Manager overload, org design misalignment, decision velocity loss, span-of-control drift', 'STAR model scores, span of control, org layer ratio, engagement by team'],
        ['A', 'Acquire', 'Deliver talent at the right place, on time, at the right cost', 'Late hiring, clearance lead time underestimation, TOC bottlenecks, programme milestone risk', '"Ahead of Ready" status, lead time vs target, requisition age, pipeline coverage ratio'],
        ['D', 'Develop', 'Build skills and empower leaders to sustain growth', 'Competency gaps, succession voids, L&D under-investment, promotion pipeline blockages', 'Competency match %, succession depth, training hours, promotion rate by grade'],
    ]
)

heading('4.3 Theory of Constraints — Talent as a Product', 2)
body(
    'The Theory of Constraints (TOC), applied to talent, treats the workforce pipeline as '
    'a production system. The product is a qualified, cleared, and productive person deployed '
    'to the right programme. Like any production system, throughput is limited by the binding '
    'constraint — the one step in the process that limits output regardless of how efficiently '
    'everything else runs.'
)
body(
    'In defence and aerospace, the most common binding constraints are: (1) security clearance '
    'lead time — DV clearance averages 12-18 months; (2) specialist skill scarcity — '
    'R&M Engineering and Mission Systems expertise is a finite market; (3) manager approval '
    'chain delays — internal bureaucracy slowing the hiring process below the constraint. '
    'The platform identifies the current binding constraint by measuring lead times at each '
    'stage and surfacing the one with the largest gap between actual and target.'
)
simple_table(
    ['TOC Concept', 'Talent Application', 'Platform Implementation'],
    [
        ['Identify the constraint', 'Which step in H2R slows throughput most?', 'Process mining on H2R event log — lead time by stage'],
        ['Exploit the constraint', 'Maximise throughput of the bottleneck step', 'Priority actions on constrained roles; cleared contractor bridge'],
        ['Subordinate to the constraint', 'Align other processes to support the constraint', 'Programme teams forecast demand earlier; clearance initiated before requisition'],
        ['Elevate the constraint', 'Invest to break the constraint entirely', 'Clearance coordinator; pre-cleared candidate pipeline; internal redeployment priority'],
        ['Prevent inertia — repeat', 'Monitor to catch the next constraint as it emerges', 'Platform KPIs refresh weekly; TOC dashboard on Acquire section'],
    ]
)

heading('4.4 "Ahead of Ready" — Lead Time as the Strategic Variable', 2)
body(
    'Talent must be deployed before demand is urgent. This is the "Ahead of Ready" principle. '
    'In a conventional HR model, a vacancy is created when someone leaves or when a project '
    'ramps up. In the Workforce Intelligence Platform model, the demand signal is the programme '
    'plan — milestones, phase gates, and headcount ramps that are known months or years in '
    'advance. The platform compares current supply to future demand at each lead-time horizon '
    'and flags where supply will be insufficient before the need is urgent.'
)
body(
    'Lead time is the total time from "need identified" to "qualified, cleared, and productive '
    'person in seat." For DV-cleared roles, this is 14-18 months. For SC-cleared roles, '
    '9-12 months. For uncleared professionals, 3-6 months. If a milestone is 10 months away '
    'and you need a DV-cleared role filled, you are already behind. The platform makes this '
    'visible and actionable before the milestone is missed.'
)

heading('4.5 Donella Meadows — Working On the System, Not In It', 2)
body(
    'Donella Meadows\' Thinking in Systems framework is the analytical lens for the platform. '
    'Workforce is modelled as a stock (the current pool of qualified, engaged, productive people) '
    'managed through four flows: inflows via Acquire (hiring), development via Develop (upskilling '
    'and succession), stabilisation via Retain (reducing outflows), and structural optimisation '
    'via Optimize (org design). The platform sits outside the system looking in — working ON '
    'the system, not IN it.'
)
body(
    'This philosophy has a specific implication for how the platform presents information. '
    'A dashboard that shows what is happening (lagging indicators) is working IN the system. '
    'A platform that explains why something is happening (driving indicators) and tells you '
    'what to do about it before it gets worse (leading indicators) is working ON the system. '
    'Every screen in the platform is designed to surface driving indicators first, with '
    'lagging indicators as confirmation, not as the primary signal.'
)
simple_table(
    ['Indicator Type', 'Definition', 'Platform Example', 'Action Value'],
    [
        ['Driving indicator', 'Root cause — the thing that causes the problem', 'Compa-ratio below 0.90 for 6+ months', 'HIGH — act now to prevent departure'],
        ['Leading indicator', 'Early signal — the problem is developing', 'Flight risk score above 60%; time-in-role > 4 years with no promotion', 'MEDIUM — monitor and prepare intervention'],
        ['Lagging indicator', 'Consequence — the problem has already occurred', 'Attrition rate this quarter; vacancy count', 'LOW — confirm what happened, calibrate models'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. COMPETENCY FRAMEWORK BACKBONE
# ══════════════════════════════════════════════════════════════════════════════
heading('5. Competency Framework Backbone', 1)
body(
    'The competency intelligence layer is grounded in six authoritative external frameworks. '
    'These provide the semantic backbone — occupational language, competency definitions, '
    'task inventories, and proficiency level indicators used to automatically generate '
    'competency models and UGESP-defensible legal documentation.'
)
simple_table(
    ['Framework', 'Source', 'Primary Contribution', 'Access Method'],
    [
        ['O*NET', 'US DOL / O*NET Resource Center', '900+ occupations, task inventories, KSAOs, work context', 'Web Services API (X-API-Key)'],
        ['NICE Framework v2.2', 'NIST / CISA (April 2025)', '30+ cybersecurity work roles, TKS statements, competency areas', 'JSON / XLSX download'],
        ['INCOSE SECF', 'INCOSE (July 2018)', '36 Systems Engineering competencies x 5 proficiency levels x behavioural indicators', 'Parsed from published document'],
        ['R&M BoK', 'DoW / OSD (December 2025)', 'Defense acquisition engineering lifecycle competencies by phase and functional area', 'Parsed from public release'],
        ['AFECD', 'USAF / AFPC', 'Air Force Specialty Codes, skill levels, duty descriptions (Phase 2)', 'Phase 2 — pending access'],
        ['SIOP / UGESP Methodology', 'SHRM-SIOP (2024)', 'Competency modelling methodology, legal defensibility, content validity framework', 'Reference methodology'],
        ["Bloom's Taxonomy", 'Educational literature', 'Cognitive complexity mapping for proficiency levels 1-5', 'Reference model'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. STAKEHOLDER & USER ROLES — THREE-TIER OWNERSHIP MODEL
# ══════════════════════════════════════════════════════════════════════════════
heading('6. Stakeholder & User Roles — Three-Tier Ownership Model', 1)
body(
    'Version 2.0 introduces a three-tier ownership model that defines who sees what and '
    'who owns which actions. The ownership model is not about data access control alone — '
    'it defines accountability for each action plan generated by the platform. Every action '
    'has a primary owner who is responsible for moving it to resolution.'
)
simple_table(
    ['Ownership Tier', 'Role', 'What They See', 'What They Own'],
    [
        ['Action Owner', 'Line Manager / Programme Manager', 'Their team\'s flight risk, individual action plans, upcoming milestones requiring talent, direct report competency gaps', 'Day-to-day retention conversations, development plans, programme resourcing requests'],
        ['Business Partner', 'HR Business Partner', 'Department-level ROAD health, all action plans in their portfolio, H2R case risk scores, workforce analytics for their client group', 'Compensation reviews, succession facilitation, talent acquisition briefs, escalation to Leader tier'],
        ['Leader', 'HR VP / CPO / Business Leadership', 'Enterprise-wide ROAD health, programme delivery risk summary, strategic capability gaps, return on impact across portfolio', 'Strategic workforce decisions, investment approvals, org design changes, board-level reporting'],
    ]
)

body(
    'The platform presents each ownership view as a distinct "lens" — same underlying data, '
    'different level of aggregation, different action vocabulary. Line Managers see individual '
    'names and specific conversations. HR Business Partners see department patterns and '
    'portfolio actions. Leaders see trends, risks, and investment decisions.'
)
simple_table(
    ['Role', 'User Type', 'Platform Access', 'Planning Horizon', 'Visualisation'],
    [
        ['Chief People Officer', 'Leader', 'Enterprise ROAD dashboard, strategic capability gaps, AI agent, full read', 'Strategic (3-5yr)', 'Leader view + AI agent'],
        ['HR Vice President', 'Leader', 'Full platform, enterprise-wide analytics, Leader ownership view', 'All horizons', 'Leader view + AI agent'],
        ['IT Vice President', 'Executive', 'Architecture dashboard, infrastructure metrics, Phase 2 roadmap', 'Execution + Architecture', 'IT ops view'],
        ['Programme Director', 'Consumer', 'Programme delivery risk tab — talent supply vs demand for their programmes', 'Tactical (0-12mo)', 'Programme risk view (read-only)'],
        ['HR Business Partner', 'Business Partner', 'Department ROAD health, all action plans in portfolio, AI agent, H2R case risk', 'Operational + Tactical', 'Business Partner view'],
        ['Workforce Planning Analyst', 'Power User', 'Strategic WFP, competency intelligence, cascade engine, scenario modelling', 'Strategic + Operational', 'Full analytical view'],
        ['TA Recruiter / Coordinator', 'Action Owner', 'Active case risk scores, breach alerts, lead time dashboard, pipeline coverage', 'Tactical + Execution', 'Acquire view'],
        ['HR Operations Manager', 'Action Owner', 'Full process mining, compliance monitoring, H2R execution KPIs', 'Operational + Execution', 'Execution view'],
        ['Line Manager', 'Action Owner', 'Own team action plans, direct report flight risk, competency gaps, open cases', 'Tactical', 'Action Owner view (read-only)'],
        ['Competency SME', 'Specialist', 'Competency model builder, framework backbone viewer, UGESP docs', 'Operational', 'Competency intelligence view'],
        ['Compliance Officer', 'Audit', 'Compliance sentinel, breach log, UGESP documentation audit trail', 'All horizons', 'Compliance view'],
        ['Data / Analytics Analyst', 'Technical', 'Full Databricks SQL, semantic layer, notebooks (anonymised PII)', 'All horizons', 'Databricks + Tableau'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading('7. Functional Requirements', 1)
body(
    'Requirements are prioritised using MoSCoW. Requirements prefixed FR-H2R are inherited '
    'from H2R-SRNA-001 v1.0 unchanged. Requirements prefixed FR-WIP-V1 were introduced in '
    'WIP-SRNA-001 v1.0. Requirements prefixed FR-WIP-V2 are new to this version.'
)

heading('7.1 Data Ingestion & Event Log', 2)
body('FR-D01 through FR-D09 from H2R-SRNA-001 are carried forward. Key v2.0 additions:')
req_table([
    ['FR-WIP-V2-D01', 'The system shall support a Phase 1 flat-file data architecture: SAP data exported as CSV/JSON, stored locally, pre-processed into vector stores and analytical datasets without requiring live SAP API connectivity. The flat-file schema SHALL be designed as the data requirements specification for Phase 2 IT integration.', 'Must', 'Phase 1 must work without live SAP integration to enable rapid demonstration and proof of concept. Flat file schema becomes the IT interface specification.'],
    ['FR-WIP-V2-D02', 'The system shall ingest synthetic workforce data (employees, programmes, attrition history, H2R action log) as the demonstration dataset for Phase 1. The synthetic data schema shall match the target production schema field-for-field.', 'Must', 'Demo requirement: platform must be demonstrable without access to production HR data.'],
    ['FR-WIP-V2-D03', 'The Phase 2 data architecture shall integrate with SAP SuccessFactors via OData API and SAP HCM/ECC for employee master data, position data, and H2R event logs. Phase 2 shall target nightly batch ingestion. Phase 3 shall evaluate event-driven ingestion.', 'Must', 'IT integration requirement: Phase 2 connects to live SAP. This document is the data specification.'],
])

heading('7.2 SAP Action Log Integration (H2R Event Log)', 2)
body(
    'The H2R action log is the primary process data source. In Phase 1 it is represented as '
    'a synthetic flat file matching the SAP SuccessFactors event log schema. In Phase 2 it '
    'is ingested via OData API. The action log is the foundation of process mining and '
    'the Acquire ROAD pillar.'
)
req_table([
    ['FR-WIP-V2-AL01', 'The system shall ingest the H2R action log: a timestamped record of every activity completed within each hiring or onboarding case (requisition raised, job posted, interview completed, offer extended, clearance initiated, background check, etc.). Minimum fields: case_id, activity_name, timestamp, actor_role, outcome, days_from_previous_activity.', 'Must', 'Process mining requires a complete event log with activity-level timestamps. Without this, bottleneck identification is not possible.'],
    ['FR-WIP-V2-AL02', 'The action log shall capture the clearance initiation date for every case requiring security vetting, enabling lead time calculation from clearance initiation to clearance granted. This is the single most important date for "Ahead of Ready" analysis.', 'Must', 'Clearance lead time is the most common binding constraint in defence talent acquisition. Tracking it is non-negotiable.'],
    ['FR-WIP-V2-AL03', 'The system shall calculate a risk score for each active H2R case based on: days since last activity (staleness), number of activities remaining, SLA proximity, and compliance completeness. Risk score shall be 0-100 with labels: Low / Medium / High / Critical.', 'Must', 'Case risk scoring is a core H2R intelligence capability inherited from H2R-SRNA-001.'],
    ['FR-WIP-V2-AL04', 'The action log integration shall be a stated technical requirement (FR priority: Must) for the Phase 2 IT build. IT must provide OData API access to the SAP SuccessFactors Recruiting Management and Onboarding modules. This requirement is blocked on IT delivery.', 'Must', 'IT dependency: without action log API access, real-time process intelligence is not possible in production.'],
])

heading('7.3 Process Mining Engine (Inherited from H2R-SRNA-001)', 2)
body('FR-PM01 through FR-PM09 from H2R-SRNA-001 are carried forward unchanged. See H2R-SRNA-001 §4.2.')

heading('7.4 Programme Execution Risk Management (NEW in v2.0)', 2)
body(
    'This module connects workforce data to programme delivery risk. It is the primary '
    'differentiator of the v2.0 platform — making the link between HR decisions and '
    'business outcomes explicit and quantified.'
)
req_table([
    ['FR-WIP-V2-PER01', 'The system shall maintain a programme workforce risk assessment for each active programme: current headcount vs required headcount by role, "Ahead of Ready" status for each open role (comparing supply timeline to lead time + milestone date), and an overall delivery risk score (0-100).', 'Must', 'Programme delivery risk is the primary business driver for this platform. Every metric must trace to programme impact.'],
    ['FR-WIP-V2-PER02', 'The system shall calculate "Ahead of Ready" status for each open role on each programme: a role is Ahead of Ready if the estimated date of supply (hire date + onboarding + ramp) precedes the milestone date by the required margin. A role is At Risk if supply will arrive at or after the milestone date. A role is Behind if current lead time means the milestone will be missed with no intervention.', 'Must', 'Ahead of Ready is the core supply chain concept. Talent must be deployed before demand is urgent, not in response to it.'],
    ['FR-WIP-V2-PER03', 'The system shall identify the Theory of Constraints binding constraint in the talent pipeline: the single process stage with the largest gap between actual lead time and target lead time, representing the bottleneck that limits programme delivery readiness. Constraint shall be identified by programme and overall.', 'Must', 'TOC: identifying and exploiting the constraint is the most impactful action. All other improvements are secondary until the constraint is addressed.'],
    ['FR-WIP-V2-PER04', 'The system shall produce a programme milestone risk assessment showing each upcoming milestone, its date, the workforce readiness status (On Track / At Risk / Delayed), and the specific talent gap contributing to risk where applicable.', 'Must', 'Programme Directors and HR Leaders require a milestone-linked talent risk view, not a generic vacancy report.'],
    ['FR-WIP-V2-PER05', 'The system shall support lead time modelling: for each role type and clearance level, maintain a database of historical lead times (from need identified to productive in role) and compare to current open cases to identify cases tracking ahead of or behind their expected lead time.', 'Should', 'Lead time data enables proactive intervention. If a case is tracking 30% slower than the historical average, that is an early warning signal.'],
])

heading('7.5 Retain — Attrition & Retention Intelligence (NEW in v2.0)', 2)
body(
    'The Retain pillar stabilises the workforce stock by reducing regrettable outflows. '
    'v2.0 introduces two significant changes: (1) attrition is classified by programme '
    'execution impact, not by 9-box position; (2) turnover contagion is modelled at '
    'team level. These changes ensure retention actions are prioritised by business '
    'impact, not by HR convention.'
)
req_table([
    ['FR-WIP-V2-R01', 'The system shall classify all employee departures as Regrettable, Non-Regrettable, or Neutral based on programme execution impact, NOT based on 9-box position or performance rating. A departure is Regrettable if it introduces risk into programme delivery. Classification criteria: (a) employee was on critical path for a programme AND held a scarce clearance (SC or DV); OR (b) employee had key person dependency designation; OR (c) employee was a succession nominee at Ready Now or Development Pipeline level with high performance. Classification must be transparent and auditable.', 'Must', 'Regrettability classification drives resource priority for retention actions. Basing it on programme impact ensures HR investment is aligned to business risk, not HR convention. 9-box is explicitly excluded as it blends development with talent management — these are managed separately in this organisation.'],
    ['FR-WIP-V2-R02', 'The system shall calculate a flight risk score (0-100) for each employee using a composite of driving indicators: compa-ratio (pay vs peer midpoint), months since last promotion, engagement score, time-in-role, market demand for skills (approximated from discipline and clearance scarcity), and manager span-of-control. Each factor shall have a configurable weight. The score shall be recalculated with each data refresh.', 'Must', 'Flight risk scoring enables proactive intervention before departure intent is expressed. The score is a leading indicator — it rises before the person begins job searching.'],
    ['FR-WIP-V2-R03', 'The system shall produce a plain-English narrative for each high-risk employee\'s flight risk score, explaining which factors are driving the score and what each means in practical terms. Example: "Sarah Chen is likely to leave within 6 months. Her salary is 11% below her peer group\'s midpoint, she has not been promoted in 3 years, and her skills are in high demand in the current market." This narrative must be understandable by a Line Manager with no data science background.', 'Must', 'VUCA Prime — Clarity: HR Business Partners and Line Managers are not data scientists. Plain English is not optional.'],
    ['FR-WIP-V2-R04', 'The system shall model turnover contagion at the team level: when a regrettable departure occurs, the flight risk score for all direct team members shall be increased by a configurable uplift factor (default: 40% uplift, range 20-80%) for a configurable period (default: 6 months). Contagion events shall be surfaced as action plan triggers for the HR Business Partner.', 'Must', 'Research literature shows 40-80% increase in flight risk among direct teammates following a regrettable departure. Ignoring contagion systematically underestimates retention risk after any significant departure.'],
    ['FR-WIP-V2-R05', 'For each regrettable attrition event, the system shall record: departure reason category, knowledge documentation status (Yes/Partial/No), team contagion assessment, estimated replacement cost, and estimated weeks to replace. This forms the attrition intelligence record used for trend analysis and action plan generation.', 'Must', 'Attrition intelligence is the foundation of "cost of inaction" calculations in action plans. Without it, return on impact cannot be calculated.'],
])

heading('7.6 Optimize — Organisational Performance (NEW in v2.0)', 2)
body(
    'The Optimize pillar accelerates organisational performance by identifying structural '
    'and process issues that limit productivity. Version 2.0 introduces the Galbraith '
    'STAR Model as the primary diagnostic framework.'
)
req_table([
    ['FR-WIP-V2-OPT01', 'The system shall score each employee, team, and organisational unit against the Galbraith STAR Model: Strategy (workforce-strategy alignment), Structure (organisational design effectiveness), Processes (information flow and decision rights), Rewards (compensation and recognition effectiveness), People (development and succession quality). Scores are derived from existing survey and HR data, not from separate assessments.', 'Must', 'STAR Model provides a systems view of org effectiveness that goes beyond individual performance. It identifies structural constraints that individual interventions cannot fix.'],
    ['FR-WIP-V2-OPT02', 'The system shall analyse span of control for all managers: direct reports count, engagement score of their team, average flight risk in their team, and performance distribution. Flag managers below the minimum span (default: 4) as over-managed, and above the maximum span (default: 12) as stretched. Both are Optimize risks.', 'Must', 'Span of control has documented productivity impacts. Over-management creates bottlenecks and slows decisions. Excessive span reduces development quality and increases manager burnout risk.'],
    ['FR-WIP-V2-OPT03', 'The system shall analyse organisational layer distribution: count of employees at each layer (Executive, Senior Manager, Manager/Lead, Professional, Junior), and flag where the distribution is out of balance relative to organisation size and programme complexity benchmarks.', 'Should', 'Org layer imbalance is a leading indicator of decision velocity problems. Too many layers slow execution; too few layers overload senior people.'],
])

heading('7.7 Acquire — Talent Pipeline Management and Lead Time (NEW in v2.0)', 2)
body(
    'The Acquire pillar delivers talent at the right place, on time, at the right cost. '
    'Version 2.0 extends the inherited H2R process intelligence with "Ahead of Ready" '
    'lead time analysis and TOC constraint identification.'
)
req_table([
    ['FR-WIP-V2-ACQ01', '"Ahead of Ready" status shall be displayed for every open requisition: the required completion date (derived from programme milestone), the expected completion date (based on current progress and historical lead time), the gap in days/weeks, and the risk colour (Green/Amber/Red). Green means supply arrives before demand is urgent. Red means supply will arrive after the milestone date without intervention.', 'Must', '"Ahead of Ready" is the core supply chain concept for this platform. It transforms talent acquisition from reactive gap-filling to proactive supply chain management.'],
    ['FR-WIP-V2-ACQ02', 'The system shall maintain lead time benchmarks by role type and clearance level: Developed Vetting (DV) cleared roles, Security Check (SC) cleared roles, Baseline cleared roles, and Uncleared professionals. Benchmarks shall be updated from historical actuals each quarter. Lead time = time from requisition raised to person in seat and assessed as productive.', 'Must', 'Lead time benchmarks are the reference for "Ahead of Ready" calculation. Without them, the platform cannot determine whether a case is on track or behind.'],
    ['FR-WIP-V2-ACQ03', 'The system shall identify and surface the binding constraint in the talent pipeline using the Theory of Constraints methodology: measure actual lead time at each H2R process stage, compare to target, identify the stage with the largest gap, and surface it as the Constraint with recommended exploit/subordinate/elevate actions.', 'Must', 'TOC constraint identification directs management attention to the one action that will have the most impact on throughput. Without it, organisations improve the non-constraints and wonder why throughput does not improve.'],
    ['FR-WIP-V2-ACQ04', 'For roles requiring SC or DV clearance, the system shall track clearance initiation date separately from requisition date, and shall flag any case where clearance has not been initiated within 10 days of requisition raising. Late clearance initiation is the most common cause of "Behind" status in cleared role pipelines.', 'Must', 'Clearance initiation delay is an exploitable constraint — it can be fixed with process improvement alone, without structural change. It is the most common source of avoidable delay in defence talent acquisition.'],
])

heading('7.8 Develop — Capability Building (Inherited from H2R-SRNA-001 + Extended)', 2)
body('FR-WIP-V1-CI01 through FR-WIP-V1-CI07 from v1.0 are carried forward. Key v2.0 additions:')
req_table([
    ['FR-WIP-V2-DEV01', 'The system shall calculate a competency match percentage for each employee: the proportion of their current role\'s competency requirements met by their assessed proficiency profile. A match below 70% shall trigger a Develop action plan. A match below 50% shall trigger a Critical priority Develop action plan.', 'Must', 'Competency match drives targeted development investment. Without it, L&D investment is distributed by request, not by need.'],
    ['FR-WIP-V2-DEV02', 'The system shall assess succession depth for each key role: Ready Now (successor can take the role immediately), Development Pipeline 12-18 months (successor needs development time), Development Pipeline 24-36 months, and Not Nominated (no identified successor). Roles with key_person_dependency=True and succession_status=Not Nominated are flagged as Critical succession voids.', 'Must', 'Succession voids are "Ahead of Ready" failures in the internal talent supply chain. A key person departure with no successor is equivalent to a stock-out in manufacturing.'],
])

heading('7.9 Proactive Action Plans with Return on Impact (NEW in v2.0)', 2)
body(
    'Action plans are the primary output of the platform. Every action plan is generated '
    'from driving indicators — the root causes of workforce risk — not from lagging '
    'indicators. Every action plan includes a return on impact calculation: the cost of '
    'taking the action compared to the cost of taking no action.'
)
req_table([
    ['FR-WIP-V2-AP01', 'The system shall generate proactive action plans from driving indicator thresholds across all four ROAD pillars. An action plan shall include: (a) plain-English title; (b) driving indicator narrative explaining why this risk exists; (c) impact narrative explaining what happens if no action is taken; (d) specific recommended actions (numbered steps, not generic advice); (e) primary owner (Action Owner / Business Partner / Leader); (f) ROAD pillar tag; (g) priority level (Critical / High / Medium); (h) return on impact calculation.', 'Must', 'Action plans are the platform\'s primary value delivery. A risk score without a recommended action is a lagging indicator dashboard. The platform must tell people what to do, not just what is wrong.'],
    ['FR-WIP-V2-AP02', 'The return on impact calculation for each action plan shall include: (a) cost of the recommended action (estimated management time, compensation adjustment, training investment, recruitment cost, etc.); (b) cost of inaction (replacement cost, programme delay cost, contagion multiplier, contractor premium, etc.); (c) return ratio (cost of inaction / cost of action). All costs shall be in GBP. Replacement cost default: 140% of annual salary (configurable). Programme delay cost: configurable per programme by HR VP.', 'Must', 'Return on impact makes the business case for every intervention. Leaders and Finance need to see the financial logic before approving expenditure. This is also the platform\'s strongest differentiator from conventional HR dashboards.'],
    ['FR-WIP-V2-AP03', 'Action plans shall be filtered and presented differently for each ownership tier. Action Owner view: individual names, specific conversations, immediate next steps. Business Partner view: portfolio of actions across their client group, grouped by ROAD pillar and priority. Leader view: aggregated return on impact across portfolio, strategic decisions required, investment approvals.', 'Must', 'Ownership tier filtering is a core platform feature. Showing a Line Manager enterprise-wide trends is noise. Showing a CPO individual names is a GDPR risk.'],
    ['FR-WIP-V2-AP04', 'Action plan status shall be tracked through a defined lifecycle: Generated (system-created) → Acknowledged (owner has seen it) → In Progress (action taken) → Resolved (risk mitigated) → Escalated (moved to higher ownership tier). Status updates shall be recorded for audit trail.', 'Should', 'Lifecycle tracking enables measurement of platform effectiveness. If action plans are generated but not acknowledged, the platform is not being used as intended.'],
])

heading('7.10 Predictive Analytics (Inherited from H2R-SRNA-001 + Extended)', 2)
body('FR-PA01 through FR-PA07 from H2R-SRNA-001 are carried forward. FR-WIP-V1-PA01 through PA03 from v1.0 are carried forward. Key v2.0 addition:')
req_table([
    ['FR-WIP-V2-PA01', 'All ML model outputs shall be accompanied by a plain-English explanation of the three to five factors driving the score for each individual. The explanation shall be generated using SHAP values and translated into natural language by the AI agent. Technical ML outputs shall not be surfaced to HR users without plain-English translation.', 'Must', 'VUCA Prime — Clarity: HR Business Partners cannot act on a number without understanding why it is that number. SHAP explainability + plain English translation is a non-negotiable UX requirement.'],
])

heading('7.11 Ontology & Knowledge Graph (Inherited from H2R-SRNA-001 + Extended)', 2)
body('FR-WIP-V1-OG01 through OG10 from v1.0 are carried forward. See v1.0 §6.4 for full definitions.')

heading('7.12 Strategic Workforce Planning (Inherited from v1.0)', 2)
body('FR-WIP-V1-SWP01 through SWP06 from v1.0 are carried forward. See v1.0 §6.7.')

heading('7.13 Operational Workforce Planning (Inherited from v1.0)', 2)
body('FR-WIP-V1-OWP01 through OWP05 from v1.0 are carried forward. See v1.0 §6.8.')

heading('7.14 Tactical Workforce Planning (Inherited from v1.0)', 2)
body('FR-WIP-V1-TWP01 through TWP05 from v1.0 are carried forward. See v1.0 §6.9.')

heading('7.15 Competency Intelligence (Inherited from v1.0)', 2)
body('FR-WIP-V1-CI01 through CI07 from v1.0 are carried forward. See v1.0 §6.10.')

heading('7.16 Cascade Engine (Inherited from v1.0)', 2)
body('FR-WIP-V1-CE01 through CE05 from v1.0 are carried forward. See v1.0 §6.11.')

heading('7.17 UGESP Compliance Engine (Inherited from v1.0)', 2)
body('FR-WIP-V1-UC01 through UC04 from v1.0 are carried forward. See v1.0 §6.12.')

heading('7.18 AI Agent (Inherited + Extended)', 2)
body('FR-H2R-AI01 through AI08 and FR-WIP-V1-AI01 through AI04 are carried forward. Key v2.0 additions:')
req_table([
    ['FR-WIP-V2-AI01', 'The AI agent shall answer questions grounded in real-time workforce data for the user\'s ownership tier: an Action Owner asking "Who on my team is most at risk of leaving?" receives individual-level data. A Leader asking the same question receives a programme-level risk summary. The agent shall enforce ownership-tier data scoping automatically.', 'Must', 'Ownership tier-aware AI responses prevent both information overload and PII exposure. The agent is not a raw data query tool — it is an intelligence interface with appropriate access controls.'],
    ['FR-WIP-V2-AI02', 'All AI agent interactions shall strip employee PII (full name, grade, salary) before transmitting to the Claude API. The API receives anonymised, aggregated, or pseudonymised data only. PII is re-injected by the platform UI from local data after the API response is received.', 'Must', 'Data protection requirement: employee PII must not be sent to third-party AI APIs without explicit data processing agreements and employee consent. PII stripping is the technical control.'],
])

# ══════════════════════════════════════════════════════════════════════════════
# 8. NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading('8. Non-Functional Requirements', 1)
body('All NFR sections from H2R-SRNA-001 (Performance, Security, Availability, Scalability, Data Residency) are carried forward. Key v2.0 additions:')
req_table([
    ['NFR-WIP-V2-P01', 'Programme delivery risk assessment (all programmes, all milestones, ahead-of-ready status) shall refresh within 5 minutes of data pipeline completion. Stale data must be flagged with a timestamp warning.', 'Must', 'Programme decisions are time-sensitive. A risk view that is 24 hours stale can miss a developing crisis.'],
    ['NFR-WIP-V2-S01', 'Framework backbone data (O*NET, NICE, INCOSE SECF, R&M BoK) shall be stored in a read-only immutable Bronze layer. No process shall be able to modify framework source data.', 'Must', 'Data integrity: framework backbone is the authoritative source for competency definitions. Modification would undermine UGESP defensibility.'],
    ['NFR-WIP-V2-P02', 'The Phase 1 static demo deployment (GitHub Pages) shall load all data and render the full dashboard within 3 seconds on a standard corporate laptop browser. JSON data files shall be pre-processed and minimised for performance.', 'Must', 'Demo requirement: a slow demo loses executive attention. 3 seconds is the maximum acceptable load time for a demonstration context.'],
    ['NFR-WIP-V2-SC01', 'The flat-file schema used in Phase 1 shall be maintained as the canonical data requirements specification for Phase 2. IT integration in Phase 2 shall populate exactly the same fields as the flat-file schema. No field shall exist in the flat file that is not deliverable from SAP in Phase 2.', 'Must', 'Architecture continuity: Phase 1 flat files are the data requirements spec. This ensures Phase 2 IT integration does not require platform redesign.'],
])

# ══════════════════════════════════════════════════════════════════════════════
# 9. NOTIONAL ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading('9. Notional Architecture', 1)
heading('9.1 Architecture Overview', 2)
body(
    'The Workforce Intelligence Platform is a nine-layer stack extending the H2R-SRNA-001 '
    'seven-layer architecture. Phase 1 uses flat files and vector stores. Phase 2 connects '
    'to live SAP via OData API. The flat-file schema in Phase 1 is the IT data requirements '
    'specification for Phase 2.'
)
simple_table(
    ['Phase', 'Data Architecture', 'Timeframe', 'SAP Connection'],
    [
        ['Phase 1 — Demo & Proof of Concept', 'CSV/JSON flat files + ChromaDB vector store + static GitHub Pages deployment', 'Months 1-3', 'No live SAP connection — CSV export from SAP or synthetic data'],
        ['Phase 2 — IT Integration', 'SAP OData API + Databricks Delta Lake (Bronze/Silver/Gold) + Unity Catalog + near-real-time pipeline', 'Months 4-12', 'Live SAP SuccessFactors and HCM/ECC connection via OData'],
        ['Phase 3 — Event-Driven', 'Event-driven ingestion (SAP Change Data Capture), real-time ROAD action routing, self-updating cascade engine', 'Year 2+', 'Event-driven SAP integration via ADF/Kafka'],
    ]
)

heading('9.2 Layer Definitions', 2)
simple_table(
    ['Layer', 'Name', 'Components', 'Phase 1 State', 'Phase 2 State'],
    [
        ['L0', 'Source Systems', 'SAP SuccessFactors, SAP HCM/ECC, SAP Payroll, O*NET API, NICE JSON, INCOSE SECF, R&M BoK', 'CSV export from SAP; framework files downloaded', 'Live OData API connection'],
        ['L1', 'Ingestion', 'CSV loader, framework parsers, O*NET API connector', 'Python script loading flat files', 'ADF / AWS Glue pipeline; OData connector'],
        ['L2', 'Lakehouse', 'Databricks Delta Lake (Bronze/Silver/Gold), Unity Catalog, MLflow', 'Local file system (Phase 1 demo)', 'Full Databricks deployment'],
        ['L3', 'Semantic / Ontology', 'Semantic layer, HR ontology (OWL), Knowledge Graph (Neo4j), ChromaDB vector store', 'ChromaDB local; JSON semantic layer', 'Full Neo4j + Databricks semantic layer'],
        ['L4', 'Analytics', 'PM4Py process mining, ML models (RF+GBT), SHAP explainability, Claude API agent', 'Synthetic data analytics + Claude API RAG', 'Full ML pipeline on Databricks'],
        ['L5', 'Strategic Intelligence (NEW)', 'Strategic WFP engine, capability gap calculator, scenario modeller, competency model generator, UGESP doc engine', 'Competency model generator + scenario modeller', 'Full strategic intelligence layer'],
        ['L6', 'Cascade Engine (NEW)', 'Bidirectional cascade propagation, plan sync, cascade break detection', 'Manual cascade reporting', 'Automated cascade via Databricks Workflows'],
        ['L7', 'Visualisation', 'React dashboard, Tableau, AI agent chat UI, ROAD action plan views', 'Static HTML + Chart.js + fetch() JSON', 'React SPA + Tableau Server'],
        ['L8', 'Actions / Alerts', 'Email alerts, Teams webhook, compliance sentinel, ROAD action routing', 'Demo only — no live alerts', 'Teams/email integration; action routing to owner tiers'],
    ]
)

heading('9.3 Synthetic Data Schema — Phase 1 Data Dictionary', 2)
body(
    'The following table defines the canonical data schema used in Phase 1 synthetic data '
    'and specifies the SAP source field for Phase 2 integration. This table IS the data '
    'requirements specification for IT.'
)
simple_table(
    ['Field', 'Type', 'Description', 'Phase 1 Source', 'Phase 2 SAP Source'],
    [
        ['id', 'String', 'Unique employee identifier', 'Generated (EMP1001+)', 'SAP Employee ID (Infotype 0000)'],
        ['name', 'String', 'Full name (PII — masked in AI calls)', 'Synthetic', 'SAP SuccessFactors person entity'],
        ['department', 'String', 'Organisational department', 'Synthetic', 'SAP Org Unit (Infotype 0001)'],
        ['grade', 'String', 'Pay grade with label', 'Synthetic', 'SAP Pay Scale Group (Infotype 0008)'],
        ['clearance', 'Enum', 'Security clearance level (Baseline/SC/DV)', 'Synthetic', 'HRMS clearance field (custom infotype)'],
        ['flight_risk_score', 'Integer 0-100', 'Composite flight risk indicator', 'Calculated from composite formula', 'Calculated from SAP data'],
        ['compa_ratio', 'Float', 'Salary vs peer group midpoint', 'Synthetic', 'SAP Payroll / comp planning tool'],
        ['engagement_score', 'Integer 0-100', 'Latest engagement survey score', 'Synthetic', 'Engagement survey integration (Phase 2)'],
        ['would_be_regrettable', 'Boolean', 'Departure would introduce programme execution risk', 'Calculated from programme impact criteria', 'Calculated from SAP + programme data'],
        ['on_critical_path', 'Boolean', 'Employee assigned to programme critical path', 'Synthetic', 'Programme management system integration (Phase 3)'],
        ['succession_status', 'Enum', 'Succession readiness level', 'Synthetic', 'SAP SuccessFactors Succession Planning module'],
        ['star_strategy / structure / processes / rewards / people / overall', 'Integer 0-100', 'Galbraith STAR Model scores', 'Synthetic', 'Derived from engagement survey + HR data'],
        ['programs', 'Array[String]', 'Programme IDs this employee is allocated to', 'Synthetic', 'SAP Project System / PS module'],
        ['program_allocation', 'Dict[String, Integer]', 'Programme allocation percentages', 'Synthetic', 'SAP PS / project costing'],
        ['key_person_dependency', 'Boolean', 'Flagged as key person (departure = critical risk)', 'Synthetic', 'HR BP designation in SuccessFactors'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. VISUALISATION ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading('10. Visualisation Architecture', 1)
body('Visualisation requirements from H2R-SRNA-001 §8 are carried forward. Key v2.0 additions:')
simple_table(
    ['Screen / View', 'Primary Audience', 'Key Content', 'Ownership Tier'],
    [
        ['Overview', 'All tiers', 'VUCA Prime status strip, ROAD health bars, delivery risk snapshot, top priority actions', 'All'],
        ['Programme Delivery Risk', 'Leader + Business Partner', '"Ahead of Ready" supply vs demand bars, TOC constraint, milestone risk table', 'Leader + BP'],
        ['Retain', 'Business Partner + Action Owner', 'Flight risk ranking, attrition classification, contagion events, regrettable departure trend', 'BP + AO'],
        ['Optimize', 'Business Partner + Leader', 'STAR Model scores, span of control distribution, org layer analysis', 'BP + Leader'],
        ['Acquire', 'Action Owner + Business Partner', 'H2R case risk scores, lead time vs benchmark, pipeline coverage, clearance status', 'AO + BP'],
        ['Develop', 'Business Partner + Action Owner', 'Competency match distribution, succession depth, training investment vs gap closure', 'BP + AO'],
        ['Action Plans', 'All tiers (filtered)', 'Full action plan cards with driving indicators, impact, recommended steps, and return on impact', 'All (filtered by tier)'],
        ['People Data', 'Business Partner', 'Searchable, filterable employee table with key workforce attributes', 'BP only'],
        ['AI Agent', 'All tiers', 'Natural language workforce questions grounded in real-time data, tier-aware responses', 'All (tier-scoped)'],
        ['Architecture', 'IT VP + Leader', 'Nine-layer platform architecture, phase 1/2/3 roadmap, technology decisions', 'IT + Leader'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. TECHNOLOGY DECISIONS REQUIRED
# ══════════════════════════════════════════════════════════════════════════════
heading('11. Technology Decisions Required', 1)
body('Technology decisions TD-01 through TD-15 from v1.0 are carried forward. Key v2.0 additions:')
simple_table(
    ['Decision ID', 'Decision Required', 'Default / Recommendation', 'Phase 1 Blocking?', 'Owner'],
    [
        ['TD-16', 'SAP SuccessFactors OData API access for H2R action log (Phase 2)', 'IT to provide API credentials and endpoint documentation', 'NO — Phase 1 uses flat files. BLOCKS Phase 2.', 'IT VP'],
        ['TD-17', 'SAP HCM/ECC employee master data API or CSV export format (Phase 2)', 'IT to provide SAP HR Admin export spec or OData connection', 'NO — Phase 1 uses flat files. BLOCKS Phase 2.', 'IT VP'],
        ['TD-18', 'Engagement survey platform integration (Phase 2)', 'IT to identify current engagement survey platform and data export format', 'NO — Phase 1 uses synthetic engagement scores', 'HR VP'],
        ['TD-19', 'Anthropic Data Processing Agreement (DPA) for Claude API in production', 'Legal to review Anthropic DPA terms before production AI agent deployment', 'NO — demo uses Claude API without live PII. BLOCKS Phase 2 production.', 'Legal / CISO'],
        ['TD-20', 'Programme critical path data source (which system tracks employee-to-programme-critical-path allocation?)', 'Likely SAP PS or project management tool — IT to confirm', 'NO — Phase 1 uses synthetic data. BLOCKS accurate "Ahead of Ready" in Phase 2.', 'IT VP + Programme Directors'],
        ['TD-21', 'Security clearance status data source (which system tracks clearance level per employee?)', 'Likely custom SAP infotype or HRMS — IT to confirm field mapping', 'NO — Phase 1 uses synthetic data. Critical for DV/SC lead time analysis in Phase 2.', 'IT VP + CISO'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 12. ASSUMPTIONS & CONSTRAINTS
# ══════════════════════════════════════════════════════════════════════════════
heading('12. Assumptions & Constraints', 1)
heading('12.1 Assumptions', 2)
body('Assumptions A01-A17 from v1.0 are carried forward. v2.0 additions:')
bullet('A18: The synthetic data schema (95 employees, 5 programmes, 22 attrition records, 40 H2R action log records) accurately represents the structure of AeroDefend Group\'s production data. The schema was designed with this assumption and field-by-field mapping to SAP infotypes.')
bullet('A19: Employee security clearance level is tracked in a structured field in SAP or HRMS. Without structured clearance data, "Ahead of Ready" analysis for cleared roles is not possible.')
bullet('A20: Programme critical path allocation is tracked at the employee level in a system accessible to IT. If this data exists only in programme managers\' spreadsheets, Phase 2 integration requires a data governance decision to centralise it.')
bullet('A21: The ROAD framework (Retain, Optimize, Acquire, Develop) is acceptable to HR leadership as the primary strategic framework for workforce action. v2.0 is built around ROAD as the navigation and action backbone.')
bullet('A22: Regrettability classification by programme execution impact (not 9-box) is acceptable to HR leadership and will not conflict with existing performance management processes.')
bullet('A23: The three-tier ownership model (Action Owner / Business Partner / Leader) accurately reflects the decision rights model in AeroDefend Group. Platform access scoping will be implemented according to this model.')
bullet('A24: The "Ahead of Ready" lead time benchmarks in the Phase 1 synthetic data (DV: 14-18mo, SC: 9-12mo, Baseline: 3-6mo) are representative of AeroDefend Group\'s actual lead times. These will be calibrated from historical SAP data in Phase 2.')

heading('12.2 Constraints', 2)
body('Constraints C01-C09 from v1.0 are carried forward. v2.0 additions:')
bullet('C10: The platform is analytics-only — no write-back to SAP, no case creation, no approval routing. Case management stays in SAP in all phases. The platform generates action recommendations; humans execute them in SAP.')
bullet('C11: 9-box grid is explicitly excluded from this platform. Attrition regrettability and talent tiering are based on programme execution impact, not on performance/potential grids. This is a design constraint, not a technical constraint.')
bullet('C12: The flat-file schema used in Phase 1 cannot be changed without reviewing the impact on Phase 2 IT integration planning. Schema changes require agreement between HR Transformation and IT Architecture.')

# ══════════════════════════════════════════════════════════════════════════════
# 13. OPEN QUESTIONS REGISTER
# ══════════════════════════════════════════════════════════════════════════════
heading('13. Open Questions Register', 1)
body('Open questions from v1.0 are carried forward. v2.0 additions:')
simple_table(
    ['Ref', 'Question', 'Owner', 'Priority', 'Status'],
    [
        ['V2-SAP-01', 'What SAP infotype or field stores employee security clearance level? This is required for "Ahead of Ready" analysis for cleared roles in Phase 2.', 'IT VP', 'CRITICAL', 'Open'],
        ['V2-SAP-02', 'What system tracks which employees are on the critical path for each programme? Is this in SAP PS, a project management tool, or Programme Directors\' own records?', 'IT VP + Programme Directors', 'CRITICAL', 'Open'],
        ['V2-SAP-03', 'What is the H2R action log format in SAP SuccessFactors — is it available via OData API? What activities are logged? Is there a timestamp per activity?', 'IT VP', 'CRITICAL', 'Open'],
        ['V2-ROAD-01', 'Has HR VP and CPO reviewed the ROAD framework? Are the four pillars (Retain, Optimize, Acquire, Develop) and their definitions acceptable as the platform\'s strategic architecture?', 'HR VP / CPO', 'HIGH', 'Open'],
        ['V2-ROAD-02', 'Is the regrettability classification logic acceptable (programme execution impact, not 9-box)? Who has approval authority to designate an employee as "would be regrettable"?', 'HR VP', 'HIGH', 'Open'],
        ['V2-ENGAGE-01', 'What engagement survey platform does AeroDefend Group use? Is data exportable in structured format? How frequently are surveys run?', 'HR VP', 'HIGH', 'Open'],
        ['V2-OWNER-01', 'Is the three-tier ownership model (Action Owner / Business Partner / Leader) aligned to AeroDefend Group\'s HR operating model? Who has authority to update the ownership designation for each action?', 'HR VP', 'MEDIUM', 'Open'],
        ['V2-LEGAL-01', 'Has Legal reviewed the turnover contagion model and its plain-English output? Is it acceptable to surface in writing that specific team members have elevated flight risk as a result of a colleague\'s departure?', 'Legal / HR VP', 'HIGH', 'Open'],
    ]
)

body('\nNext Steps — v2.0 (carried forward)', bold=True)
bullet('1. HR VP and CPO to review the ROAD framework and VUCA Prime framing — confirm these are acceptable as the platform strategy (V2-ROAD-01).')
bullet('2. HR VP to review the three-tier ownership model — confirm it maps to AeroDefend Group\'s HR operating model (V2-OWNER-01).')
bullet('3. IT VP to answer V2-SAP-01 through V2-SAP-03 — clearance data, critical path data, and action log format. These BLOCK Phase 2 planning.')
bullet('4. Legal to review regrettability classification logic and turnover contagion output (V2-LEGAL-01).')
bullet('5. All: review open questions from v1.0 that remain open (WIP-OG-01, WIP-OG-02, WIP-SWP-01, WIP-SWP-02 are still critical).')
bullet('6. Programme Directors: validate the five rights of talent framing and the programme delivery risk view — confirm this is the right lens for executive workforce reporting.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# v3.0 NEW CONTENT
# ══════════════════════════════════════════════════════════════════════════════

heading('14. Enterprise Scale Architecture (NEW in v3.0)', 1)
body(
    'Version 3.0 introduces an enterprise-scale reference architecture representing a '
    '50,000-employee defence and aerospace organisation. This replaces the small '
    'synthetic dataset used in v1.0 and v2.0 with a representative sample of 274 '
    'employees across 12 functional organisations, 9 active programmes, and 3 portfolios. '
    'Scale factors allow every metric to be extrapolated to the full 54,500-FTE '
    'organisation while maintaining manageable data volumes for the Phase 1 demonstration.'
)

heading('14.1 Enterprise Organisational Structure', 2)
simple_table(
    ['Portfolio', 'Programmes', 'Contract Value', 'Delivery Risk'],
    [
        ['Air Systems', 'Typhoon Modernisation Block 20, Future Combat Air System (FCAS), Eurofighter E-Scan Radar Upgrade', '£1,620M', 'High (avg 58)'],
        ['Ground Defence & Missiles', 'Air Defence Command & Control, SKYSHIELD Ground Integration, SHORAD Missile Enhancement', '£1,620M', 'Medium (avg 42)'],
        ['Cyber, EW & ISR', 'Cyber Mission System, Electronic Attack Platform, ISR Data Fusion Platform', '£980M', 'High (avg 68)'],
    ]
)

heading('14.2 Functional Organisation Registry', 2)
body(
    'Each functional organisation is the supply side of the matrix. Functions own people, '
    'develop capability, and allocate resources to programmes. The scale factor converts '
    'sample employee counts to actual workforce size for extrapolated analytics.'
)
simple_table(
    ['Function', 'Code', 'Actual FTEs', 'Sample Size', 'Scale Factor', 'Market Demand'],
    [
        ['Systems Engineering', 'SE', '9,200', '46', '200x', 'Very High'],
        ['R&M Engineering', 'RM', '7,800', '39', '200x', 'High'],
        ['Software Engineering', 'SW', '6,200', '31', '200x', 'Critical'],
        ['Cybersecurity', 'CS', '3,800', '19', '200x', 'Critical'],
        ['Mission Systems', 'MS', '5,400', '27', '200x', 'Very High'],
        ['Systems Integration & Test', 'ST', '4,100', '21', '195x', 'High'],
        ['Programme Management', 'PM', '4,600', '23', '200x', 'Medium'],
        ['Quality & Safety', 'QS', '2,900', '15', '193x', 'Medium'],
        ['Logistics & Supportability', 'LS', '3,200', '16', '200x', 'Medium'],
        ['Finance & Commercial', 'FC', '2,100', '11', '191x', 'Low'],
        ['Human Resources', 'HR', '1,800', '9', '200x', 'Low'],
        ['IT & Digital Engineering', 'IT', '3,400', '17', '200x', 'High'],
    ]
)

heading('14.3 Dual-View Platform Architecture (NEW in v3.0)', 2)
body(
    'Version 3.0 introduces a Functional View / Program View toggle as the primary '
    'navigation paradigm. This directly reflects the matrix organisation model: functions '
    'are the supply side (who owns people) and programmes are the demand side (who '
    'defines the work). Users switch between views with a single toggle; a scope selector '
    'drills into a specific function or programme. The intersection is the resource '
    'allocation matrix.'
)
simple_table(
    ['View', 'Primary User', 'Scope Selector', 'Primary Analytics', 'Default Drill-In'],
    [
        ['Functional View', 'Functional VP / HR Business Partner', 'Select one of 12 functions (or All)', 'ROAD health for the function, temporal dimensions, supply vs demand, subdiscipline breakdown', 'Temporal Analysis — time in role / level / programme / lifecycle'],
        ['Program View', 'Programme Director / Portfolio Manager', 'Select one of 9 programmes (or portfolio)', 'Resource allocation, delivery risk, phase experience gaps, milestone status', 'Resource Allocation Matrix — Function x Programme'],
        ['Enterprise / AI Workspace', 'HR VP / CPO / All users', 'Enterprise-wide or scoped', 'AI agent answers any question for any scope; prompt library; action plans', 'AI Workspace (default landing view)'],
    ]
)

heading('15. Four Temporal Dimensions (NEW in v3.0)', 1)
body(
    'Version 3.0 introduces four temporal dimensions as primary analytical lenses. '
    'These dimensions answer the question "are the right people in the right place at '
    'the right time?" — not just for current state, but for trajectory. Each dimension '
    'has defined thresholds that trigger action plan recommendations.'
)

simple_table(
    ['Dimension', 'Definition', 'Optimal Range', 'Action Trigger', 'Platform Signal'],
    [
        ['Time in Role', 'How long the employee has been in their current position title', '1 – 3 years', '< 6 months (still ramping) or > 5 years (stale — rotation or career action needed)', 'Stale label + action plan if > 5yr and would be regrettable to lose'],
        ['Time in Level', 'How long the employee has been at their current pay grade / career level', '1 – 4 years', '> 4 years — overdue for promotion review; leading indicator of departure intent', 'Overdue label + retention risk flag; linked to flight risk score'],
        ['Time on Programme', 'How long the employee has been working on their primary programme', '1 – 4 years', '< 1 year (still ramping to programme knowledge) or > 5 years (rotation risk — knowledge hoarding)', 'Rotation risk label; informs programme knowledge continuity planning'],
        ['Acquisition Lifecycle Experience', 'Which acquisition phases (MSA, TMRR, EMD, P&D, O&S) the employee has worked in previously', 'At least 2 phases including current', 'Employee has no prior experience in their current programme\'s phase — phase experience gap', 'Phase gap flag (red) or phase match (green) per employee; drives Develop action plans'],
    ]
)

heading('15.1 Phase Experience Gap Analysis', 2)
body(
    'The phase experience gap is a new risk indicator introduced in v3.0. It measures '
    'whether the people assigned to a programme have prior experience working in that '
    'programme\'s current acquisition phase. A person assigned to an EMD-phase programme '
    'with no prior EMD experience represents a capability gap that formal competency '
    'scores may not capture — because the competency exists but the phase-specific '
    'judgement does not.'
)
req_table([
    ['FR-WIP-V3-TMP01', 'The system shall record and display four temporal dimensions for every employee: time_in_role_yrs, time_in_level_yrs, time_on_program_yrs, and acquisition_phases_experienced (list of phases worked in). All four shall be refreshed with each data pipeline run.', 'Must', 'Temporal dimensions are the primary new analytical capability in v3.0. They are not derivable from competency assessments and require dedicated data fields.'],
    ['FR-WIP-V3-TMP02', 'The system shall calculate an acquisition_phase_match flag for each employee assigned to a programme: True if the employee has prior experience in the programme\'s current phase; False (phase gap) if not. Phase gap shall trigger a Develop action plan with recommended knowledge transfer approach.', 'Must', 'Phase experience gap is a leading indicator of delivery quality risk. Programmes in EMD phase staffed by people with only O&S experience will exhibit predictable quality issues at design reviews.'],
    ['FR-WIP-V3-TMP03', 'The Temporal View shall display four distribution bar charts: time in role (ramping / optimal / consider rotation / stale), time in level (on track / promotion window / overdue), time on programme (new / building / deep knowledge / rotation risk), and acquisition lifecycle phase coverage (which phases the selected function\'s employees have worked in, with phase match rate for employees currently on programmes in each phase).', 'Must', 'Temporal distribution charts give functional VPs and HR Business Partners the workforce trajectory picture — not just current state, but where people are heading.'],
    ['FR-WIP-V3-TMP04', 'The system shall generate a temporal risk table listing employees with at least one elevated temporal indicator (time_in_level > 4yr OR time_in_role > 5yr OR time_on_program > 5yr OR phase_match = False), sorted by combined risk score. This table is filterable by function and subdiscipline.', 'Must', 'The temporal risk table is the primary action list for functional VPs making rotation, promotion, and development decisions.'],
])

heading('16. Resource Allocation Matrix (NEW in v3.0)', 1)
body(
    'The Resource Allocation Matrix is the intersection model connecting functional supply '
    'to programme demand. It shows, for every function-programme pair, the headcount '
    'required, the percentage of the function allocated to that programme, the average '
    'flight risk of allocated employees, and the phase experience gap count. '
    'The matrix can be toggled between three views: Headcount, Percentage of Function, '
    'and Flight Risk — allowing the user to switch between a volume view, a capacity '
    'utilisation view, and a risk view without leaving the screen.'
)
req_table([
    ['FR-WIP-V3-ALM01', 'The system shall maintain a resource allocation matrix: for every function-programme pair where the programme requires resources from that function, record headcount_required, pct_of_function (what share of the function is allocated to this programme), avg_flight_risk (average flight risk of employees on this allocation), and phase_experience_gap_count (number of employees on this allocation who lack experience in the programme\'s current phase).', 'Must', 'Resource allocation is the primary bridge between functional view and program view. Without it, functional VPs and programme directors cannot see the same data from their respective perspectives.'],
    ['FR-WIP-V3-ALM02', 'The allocation matrix shall be colour-coded by criticality of the function-programme relationship: Critical (red), High (amber), Medium (blue), Low (grey). Criticality is defined by the programme team based on whether delivery would be compromised without this function\'s contribution.', 'Must', 'Criticality colouring directs management attention to the allocations that matter most. A high-flight-risk allocation in a Critical cell is a priority action; the same risk in a Low cell is not.'],
    ['FR-WIP-V3-ALM03', 'The allocation matrix shall support three display modes toggled by the user: Headcount (FTEs required from each function per programme), Percentage of Function (what share of the function\'s total headcount is committed to this programme), and Flight Risk (average flight risk score of employees in each allocation cell).', 'Must', 'Different stakeholders need different views of the same allocation data. A functional VP wants to know capacity utilisation (% view). A programme director wants to know whether they have enough bodies (headcount view). HR BP wants to know retention risk (flight risk view).'],
    ['FR-WIP-V3-ALM04', 'The allocation matrix shall be filterable by programme (showing all functions allocated to one programme) and by function (showing all programmes that function supplies to). These filters correspond to the Program View and Functional View respectively, providing the intersection model.', 'Must', 'Filter by programme = Program View. Filter by function = Functional View. This makes the allocation matrix the single intersection point between the two views.'],
])

heading('17. AI-Integrated Prompt Library (NEW in v3.0)', 1)
body(
    'Version 3.0 integrates a structured prompt library of 280 diagnostic workforce '
    'intelligence prompts into the AI Workspace. The prompts are organised across five '
    'capability domains and are scope-aware: the same prompt answered for "Cybersecurity '
    'function" returns different data-grounded analysis than when answered for "FCAS '
    'programme" or "enterprise-wide". The prompt library replaces the need for users to '
    'know how to formulate analytical questions — it provides a curated set of questions '
    'that the platform can answer from its data.'
)
simple_table(
    ['Capability Domain', 'Prompt Count', 'Primary Users', 'Example Prompts'],
    [
        ['Diagnostic Skills', '68', 'HR BP, Workforce Planning Analyst', 'Audit our workforce data quality; assess AI readiness; prioritise use cases; build vs buy analysis'],
        ['Workforce Planning', '72', 'HR VP, CPO, Workforce Planning Analyst', 'Demand forecasting; supply modelling; scenario planning; gap analysis'],
        ['Skills & Capability', '72', 'HR BP, L&D, Competency SME', 'Skills taxonomy; capability assessment; reskilling roadmap; critical roles and succession'],
        ['Org Design', '64', 'HR VP, Business Leaders', 'Structural alternatives; spans and layers; role clarity; effectiveness measurement'],
        ['Transformation', '78', 'CHRO, HR VP, Change Lead', 'Change strategy; stakeholder mapping; communications; measurement framework'],
    ]
)

req_table([
    ['FR-WIP-V3-PL01', 'The AI Workspace shall include a prompt library panel displaying all 280 prompts organised by the five capability domains, with each domain expandable to show sub-categories and individual prompts. The panel shall include a search box to filter prompts by keyword across all domains.', 'Must', 'Prompt library discoverability is critical for adoption. HR Business Partners and Functional VPs are not data scientists — they need a curated question menu, not a blank chat interface.'],
    ['FR-WIP-V3-PL02', 'Clicking any prompt in the library shall: (1) set the prompt text in the chat input, (2) scope the response to the currently selected function or programme, and (3) send the query to the AI analytical engine immediately. The response shall be data-grounded where the synthetic (or production) data can answer the question, and framework-based where it cannot.', 'Must', 'One-click prompt execution with automatic scope context removes the need for users to know how to formulate queries or remember to set context.'],
    ['FR-WIP-V3-PL03', 'The full prompt library shall also be accessible as a standalone page with a search function and a "Run in Agent" button on every prompt. This provides a reference view for users who want to browse the full library before selecting a question.', 'Should', 'Some users prefer to browse before querying. The full-page view supports this while the inline panel supports the quick-query workflow.'],
    ['FR-WIP-V3-PL04', 'Prompt responses shall be scoped to the AI agent scope selector: if scope is set to "Systems Engineering", all responses use only Systems Engineering employee data. If scope is "Enterprise", all responses use the full dataset. Scope shall be visible to the user at all times during the AI session.', 'Must', 'Scope awareness prevents misleading analysis. A Functional VP asking about flight risk should see their function\'s data, not enterprise-wide averages that mask their specific situation.'],
])

heading('18. Updated Synthetic Data Schema — v2 (v3.0 Data Dictionary)', 1)
body(
    'The synthetic data schema was extended in v3.0 to support the new temporal '
    'dimensions, functional organisation structure, portfolio hierarchy, and resource '
    'allocation matrix. The following tables define the canonical schemas for each '
    'data asset. These schemas are the IT data requirements specification for Phase 2 '
    'SAP integration — every field must be deliverable from SAP or a connected system.'
)

heading('18.1 Employee Record Schema (employees_v2.json)', 2)
simple_table(
    ['Field', 'Type', 'New in v3?', 'Description', 'Phase 2 SAP Source'],
    [
        ['function_id', 'String', 'Yes', 'Functional organisation code (e.g. FN-SE)', 'SAP Org Unit mapped to function registry'],
        ['function_name', 'String', 'Yes', 'Full name of functional organisation', 'Derived from function_id lookup'],
        ['subdiscipline', 'String', 'Yes', 'Specific technical area within function (e.g. Mission Systems Architecture)', 'Position description / job architecture field'],
        ['time_in_role_yrs', 'Float', 'No (v2)', 'Years in current position title', 'Derived from SAP position start date'],
        ['time_in_level_yrs', 'Float', 'Yes', 'Years at current pay grade (may differ from time_in_role)', 'Derived from SAP grade change date (Infotype 0008)'],
        ['time_on_program_yrs', 'Float', 'No (v2)', 'Years on primary programme assignment', 'Derived from SAP PS project assignment date'],
        ['acquisition_phase_current', 'Enum', 'Yes', 'Current phase of primary programme (MSA/TMRR/EMD/P&D/O&S)', 'Derived from programme registry (programs_v2.json)'],
        ['acquisition_phases_experienced', 'Array', 'Yes', 'List of acquisition phases worked in historically', 'HR system field — requires new data capture process in Phase 2'],
        ['acquisition_phase_match', 'Boolean', 'Yes', 'True if employee has prior experience in current programme phase', 'Calculated from acquisition_phases_experienced vs acquisition_phase_current'],
    ]
)

heading('18.2 New Data Assets — v3.0', 2)
simple_table(
    ['Asset', 'File', 'Records', 'Description', 'Phase 2 Source'],
    [
        ['Portfolios', 'portfolios.json', '3', 'Portfolio hierarchy with programme linkage and delivery risk rollup', 'Programme management system / Portfolio Director input'],
        ['Functions', 'functions.json', '12', 'Functional organisations with subdisciplines, headcount actuals, scale factors, and ROAD health scores', 'SAP Org Management + HR function registry'],
        ['Programmes v2', 'programs_v2.json', '9', 'Programmes with portfolio_id, acquisition phase, function requirements, and criticality', 'Programme management system (SAP PS or equivalent)'],
        ['Employees v2', 'employees_v2.json', '274', 'Extended employee schema with temporal dimensions and function/subdiscipline fields', 'SAP HCM/ECC + SuccessFactors + programme system'],
        ['Allocation Matrix', 'allocation_matrix.json', '54', 'Function x Programme allocation data: headcount required, % of function, flight risk, phase gap count', 'Derived from programme requirements + employee allocation data'],
    ]
)

heading('19. v3.0 Open Questions Register', 1)
body('The following open questions are new in v3.0 and must be resolved before Phase 2 planning can proceed.')
simple_table(
    ['Ref', 'Question', 'Owner', 'Priority', 'Status'],
    [
        ['V3-FUNC-01', 'Does AeroDefend Group use a formal functional organisation model (Chief Engineer, Chief of R&M, etc.) that maps to the 12 functions defined in this document? If the names or structure differ, the function registry must be updated before Phase 2 data integration.', 'HR VP + Engineering Leadership', 'CRITICAL', 'Open'],
        ['V3-FUNC-02', 'Who is the data owner for functional organisation assignment per employee? Is this field maintained in SAP Org Management, or does it require a new field in SuccessFactors?', 'IT VP + HR VP', 'CRITICAL', 'Open'],
        ['V3-TEMP-01', 'What is the current data source for grade change date (time in level)? Is Infotype 0008 reliably populated with accurate change dates, or is this a data quality risk?', 'IT VP + HR Operations', 'HIGH', 'Open'],
        ['V3-TEMP-02', 'Acquisition phases experienced (historical) is a new data field. How will this be captured? Options: (a) derive from historical programme assignments in SAP PS; (b) self-declaration by employee; (c) manager verification. Each has different accuracy and effort tradeoffs.', 'HR VP + IT VP', 'HIGH', 'Open'],
        ['V3-PROG-01', 'What system is the authoritative source for programme phase (MSA/TMRR/EMD/P&D/O&S) and programme-to-employee assignment? Is this in SAP PS, a programme management tool (e.g. Primavera), or maintained manually?', 'IT VP + Programme Directors', 'CRITICAL', 'Open'],
        ['V3-PORT-01', 'Is the portfolio hierarchy (3 portfolios, 9 programmes) as defined in this document an accurate representation of AeroDefend Group\'s programme portfolio? Portfolio Directors must validate before Phase 2.', 'CPO + Portfolio Directors', 'HIGH', 'Open'],
        ['V3-ALLOC-01', 'Who maintains the resource allocation plan (Function x Programme headcount requirements)? Is this held in SAP PS, a resource management tool, or programme team spreadsheets? The allocation matrix cannot be automated without a structured data source.', 'IT VP + Programme Directors + Functional VPs', 'CRITICAL', 'Open'],
        ['V3-PLIB-01', 'The 280-prompt library covers Diagnostic Skills, Workforce Planning, Skills & Capability, Org Design, and Transformation. Which of the five domains should be prioritised for the Phase 2 deployment with real data? Recommend: Workforce Planning and Skills & Capability first.', 'HR VP + CPO', 'MEDIUM', 'Open'],
    ]
)

body('\nNext Steps — v3.0', bold=True)
bullet('1. Portfolio Directors: validate the portfolio and programme structure (V3-PORT-01). This is a blocking dependency for Phase 2 resource allocation matrix.')
bullet('2. Functional VPs: confirm the 12 functional organisation names and structure (V3-FUNC-01). If the org structure differs, the function registry needs updating before data integration.')
bullet('3. IT VP: confirm data sources for (a) functional organisation assignment, (b) grade change date, (c) programme phase, and (d) programme-to-employee allocation (V3-FUNC-02, V3-TEMP-01, V3-PROG-01, V3-ALLOC-01).')
bullet('4. HR VP + IT VP: decide on the approach for capturing acquisition lifecycle phase experience history (V3-TEMP-02). This is a new data field with no existing SAP source.')
bullet('5. HR VP + CPO: prioritise which of the five prompt library domains to deploy first with real data in Phase 2 (V3-PLIB-01).')
bullet('6. All v2.0 next steps remain open unless explicitly resolved. See Section 13.')

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = r'C:\Users\traft\Desktop\WIP-SRNA-001-v3.0.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
